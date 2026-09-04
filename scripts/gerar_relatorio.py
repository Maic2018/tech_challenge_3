"""Gera o Relatório de Entrega da Fase 2 do ToggleMaster em formato .docx."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Relatorio_de_Entrega_Fase2.docx")

REPO_URL = "https://github.com/Maic2018/tech_challenge_3"
VIDEO_URL = "[Inserir link do vídeo no YouTube]"
BADGE_URL = "[Opcional: link do perfil/badge Google Cloud Skills Boost]"


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    set_paragraph_spacing(p, after=8)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    set_paragraph_spacing(p, after=4)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── Capa ──
    for _ in range(4):
        doc.add_paragraph()

    for text, size, bold in [
        ("TECH CHALLENGE", 22, True),
        ("ToggleMaster", 28, True),
        ("FASE 02", 18, True),
        ("DevOps e Arquitetura Cloud", 14, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    participantes = [
        ("Nome dos Participantes – Grupo 45:", True),
        ("Juliette Engel – RM 373836 – Discord: [username]", False),
        ("Matheus Yuri Rodrigues da Silva – RM 370672 – Discord: [username]", False),
        ("Michael Jonathan Venute Viana – RM 371791 – Discord: [username]", False),
        ("Perla Coutinho Barbosa – RM 371384 – Discord: [username]", False),
        ("Rodrigo de Albuquerque Andrade – RM 373167 – Discord: [username]", False),
    ]
    for line, bold in participantes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.name = "Calibri"
        r.bold = bold

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Data:")
    r.bold = True
    r.font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("29 de junho de 2026")
    r.font.size = Pt(12)

    add_page_break(doc)

    # ── Links ──
    add_heading(doc, "Links do Projeto", level=1)
    add_body(doc, f"Repositório GitHub: {REPO_URL}")
    add_body(doc, f"Branch: feature/melhorias-scripts")
    add_body(doc, f"Vídeo de demonstração (até 20 min): {VIDEO_URL}")
    add_body(doc, f"Google Cloud Skills Boost (pontuação extra, opcional): {BADGE_URL}")

    add_page_break(doc)

    # ── Sumário ──
    add_heading(doc, "Sumário", level=1)
    for item in [
        "Resumo",
        "1. Introdução e Contexto da Fase 2",
        "2. Análise e Conteinerização (Docker)",
        "3. Execução Local da Aplicação (Docker Compose)",
        "4. Infraestrutura como Código (Terraform)",
        "5. Automação do Deploy (Scripts)",
        "6. Configuração do Cluster Kubernetes",
        "7. Orquestração e Implantação (Manifestos)",
        "8. Escalabilidade (HPA)",
        "9. Arquitetura e Data Stores",
        "10. Segurança",
        "11. Desafios Enfrentados",
        "12. Conclusão",
        "13. Referências",
    ]:
        add_body(doc, item)

    add_page_break(doc)

    # ── Resumo ──
    add_heading(doc, "Resumo", level=1)
    add_body(doc,
        "Este relatório documenta a entrega da Fase 02 do Tech Challenge POSTECH/FIAP. "
        "O ToggleMaster evoluiu de um MVP monolítico (Fase 01) para um ecossistema de cinco "
        "microsserviços conteinerizados e implantados em Kubernetes na AWS (EKS)."
    )
    add_body(doc,
        "Toda a infraestrutura de nuvem — VPC, NAT Gateway, cluster EKS, node group, RDS, "
        "ElastiCache, SQS, DynamoDB e ECR — é provisionada via Terraform (pasta infra/), "
        "utilizando a LabRole do AWS Academy. Scripts de automação na raiz do projeto "
        "(run-all.sh) orquestram o fluxo completo: verificação da conta, terraform apply, "
        "build/push das imagens, instalação do Metrics Server e Nginx Ingress, geração "
        "automática de Secrets a partir dos outputs do Terraform, deploy Kubernetes, "
        "migrations SQL e criação da API key."
    )
    add_body(doc,
        "A escalabilidade foi implementada com Horizontal Pod Autoscaler (HPA) baseado em CPU "
        "para evaluation-service e analytics-service, conforme requisito da Opção A (AWS Academy)."
    )

    # ── 1. Introdução ──
    add_heading(doc, "1. Introdução e Contexto da Fase 2", level=1)
    add_body(doc,
        "Na Fase 01, o ToggleMaster validou o conceito como monolito em EC2 + RDS. "
        "A Fase 02 exige reescrever o sistema como microsserviços distribuídos, com "
        "orquestração em Kubernetes (EKS), provisionamento de infraestrutura na AWS e "
        "demonstração de escalabilidade."
    )
    add_table(doc,
        ["Microsserviço", "Linguagem", "Responsabilidade", "Recurso de dados"],
        [
            ["auth-service", "Go", "Chaves de API e autenticação", "PostgreSQL (RDS)"],
            ["flag-service", "Python", "CRUD de feature flags", "PostgreSQL (RDS)"],
            ["targeting-service", "Python", "Regras de segmentação", "PostgreSQL (RDS)"],
            ["evaluation-service", "Go", "Hot path — decisão true/false", "Redis (ElastiCache)"],
            ["analytics-service", "Python", "Worker de eventos assíncronos", "SQS + DynamoDB"],
        ],
    )

    # ── 2. Docker ──
    add_heading(doc, "2. Análise e Conteinerização (Docker)", level=1)
    add_body(doc,
        "Conforme requisito técnico 1 da Fase 02, cada microsserviço possui Dockerfile "
        "multi-stage otimizado, com usuários non-root (Python) e binários estáticos (Go)."
    )
    add_table(doc,
        ["Serviço", "Estratégia Docker", "Porta"],
        [
            ["auth-service", "Multi-stage: golang:1.21-alpine → alpine", "8001"],
            ["flag-service", "Multi-stage: python:3.9 → python:3.9-slim + appuser", "8002"],
            ["targeting-service", "Multi-stage: python:3.9 → python:3.9-slim + appuser", "8003"],
            ["evaluation-service", "Multi-stage: golang:1.21-alpine → alpine", "8004"],
            ["analytics-service", "Multi-stage: python:3.9 → python:3.9-slim + appuser", "8005"],
        ],
    )

    # ── 3. Docker Compose ──
    add_heading(doc, "3. Execução Local da Aplicação (Docker Compose)", level=1)
    add_body(doc,
        "O arquivo docker-compose.yml na raiz sobe todo o ecossistema localmente, "
        "incluindo LocalStack para emular SQS e DynamoDB. Comando: docker compose up --build"
    )
    add_table(doc,
        ["Contêiner", "Tipo", "Porta", "Função"],
        [
            ["auth-service", "App", "8001", "Autenticação"],
            ["flag-service", "App", "8002", "CRUD de flags"],
            ["targeting-service", "App", "8003", "Segmentação"],
            ["evaluation-service", "App", "8004", "Avaliação (hot path)"],
            ["analytics-service", "App", "8005", "Worker SQS → DynamoDB"],
            ["auth-db / flag-db / targeting-db", "PostgreSQL 15", "5433–5435", "Bancos locais"],
            ["redis", "Redis 7", "6379", "Cache"],
            ["localstack", "LocalStack 3.0", "4566", "SQS + DynamoDB local"],
        ],
    )

    add_page_break(doc)

    # ── 4. Terraform ──
    add_heading(doc, "4. Infraestrutura como Código (Terraform)", level=1)
    add_body(doc,
        "Toda a infraestrutura AWS é provisionada via Terraform na pasta infra/. "
        "Um único terraform apply cria a stack completa — não há provisionamento manual "
        "pelo Console AWS. A LabRole existente do AWS Academy é referenciada via variável "
        "lab_role_arn no terraform.tfvars, conforme exigido pela Opção A do enunciado."
    )

    add_heading(doc, "4.1 Arquivos Terraform e recursos", level=2)
    add_table(doc,
        ["Arquivo", "Recursos provisionados"],
        [
            ["main.tf", "Provider AWS (região us-east-1, versão ~> 5.0)"],
            ["variables.tf", "aws_region, project_name, db_password, lab_role_arn"],
            ["terraform.tfvars", "Valores de configuração (LabRole ARN, senha DB)"],
            ["vpc.tf", "VPC 10.0.0.0/16, 2 subnets públicas, 2 privadas, IGW, SGs (EKS, RDS, Redis)"],
            ["nat.tf", "NAT Gateway, EIP, route table privada (acesso à internet dos nodes)"],
            ["eks.tf", "Launch Template (IMDS hop limit=2), cluster EKS 1.31, Managed Node Group"],
            ["rds.tf", "3× RDS PostgreSQL 15 (auth_db, flag_db, targeting_db) + DB Subnet Group"],
            ["elasticache.tf", "Cluster ElastiCache Redis 7 (cache.t3.micro)"],
            ["sqs.tf", "Fila SQS Standard togglemaster-events"],
            ["dynamodb.tf", "Tabela analytics-events (PK: event_id, PAY_PER_REQUEST)"],
            ["ecr.tf", "5 repositórios ECR (auth, flag, targeting, evaluation, analytics)"],
            ["outputs.tf", "Endpoints RDS, Redis, SQS, DynamoDB, ECR, EKS (nome, endpoint, status)"],
        ],
    )

    add_heading(doc, "4.2 Cluster EKS (eks.tf)", level=2)
    add_body(doc, "Recursos criados pelo Terraform:", bold=True)
    add_bullet(doc, "aws_eks_cluster.main — cluster togglemaster-cluster, versão Kubernetes 1.31;")
    add_bullet(doc, "role_arn = var.lab_role_arn — utiliza a LabRole do AWS Academy;")
    add_bullet(doc, "Subnets públicas e privadas em 2 AZs; endpoint público e privado habilitados;")
    add_bullet(doc, "aws_eks_node_group.main — Managed Node Group em subnets privadas;")
    add_bullet(doc, "Auto Scaling: min=1, desired=2, max=4 (conforme enunciado Fase 02);")
    add_bullet(doc, "aws_launch_template.eks_nodes — corrige IMDS http_put_response_hop_limit=2, "
                     "permitindo que pods acessem credenciais AWS via metadata da instância (SQS/DynamoDB).")

    add_heading(doc, "4.3 NAT Gateway (nat.tf)", level=2)
    add_body(doc,
        "Os nodes EKS rodam em subnets privadas. O NAT Gateway (em subnet pública) e a route "
        "table privada garantem que os pods consigam acessar a internet (pull de imagens, "
        "APIs AWS, etc.) sem expor os nodes diretamente."
    )

    add_heading(doc, "4.4 Fluxo de provisionamento", level=2)
    add_bullet(doc, "cd infra && bash 00-check-account.sh — detecta Account ID e atualiza lab_role_arn;")
    add_bullet(doc, "terraform init && terraform plan && terraform apply -auto-approve;")
    add_bullet(doc, "terraform output — exibe endpoints RDS, Redis, SQS, EKS, URLs ECR;")
    add_bullet(doc, "Tempo estimado: ~20–25 minutos (RDS e EKS são os recursos mais lentos).")

    add_heading(doc, "4.5 Checklist completo — tudo via Terraform", level=2)
    add_table(doc,
        ["Recurso Fase 02", "Provisionado", "Arquivo Terraform"],
        [
            ["VPC + Subnets + IGW", "Sim", "vpc.tf"],
            ["NAT Gateway", "Sim", "nat.tf"],
            ["Security Groups", "Sim", "vpc.tf"],
            ["Cluster EKS + LabRole", "Sim", "eks.tf"],
            ["Managed Node Group (1/2/4)", "Sim", "eks.tf"],
            ["Launch Template (IMDS fix)", "Sim", "eks.tf"],
            ["RDS PostgreSQL (×3)", "Sim", "rds.tf"],
            ["ElastiCache Redis", "Sim", "elasticache.tf"],
            ["DynamoDB", "Sim", "dynamodb.tf"],
            ["SQS", "Sim", "sqs.tf"],
            ["ECR (×5)", "Sim", "ecr.tf"],
        ],
    )

    add_page_break(doc)

    # ── 5. Scripts ──
    add_heading(doc, "5. Automação do Deploy (Scripts)", level=1)
    add_body(doc,
        "Além do Terraform, a branch feature/melhorias-scripts inclui scripts Bash que "
        "automatizam todo o pipeline de deploy, eliminando passos manuais e reduzindo erros."
    )
    add_table(doc,
        ["Script", "Função"],
        [
            ["run-all.sh", "Orquestra os 9 passos completos do zero ao ambiente validado"],
            ["infra/00-check-account.sh", "Detecta Account ID Academy e atualiza lab_role_arn no tfvars"],
            ["build-and-push.sh", "Build (--platform linux/amd64) e push das 5 imagens para ECR"],
            ["generate-secrets.sh", "Gera infra/k8s/secrets.yaml a partir dos terraform output"],
            ["deploy-k8s.sh", "Aplica manifestos K8s na ordem correta (NS → Secrets → Deployments → Ingress → HPA)"],
            ["run-migrations.sh", "Executa init.sql nos 3 RDS via pod temporário psql-migrator"],
            ["update-api-key.sh", "Cria API key real, atualiza Secret e reinicia evaluation-service"],
            ["test-fluxo-completo.sh", "Testa fluxo end-to-end via Load Balancer"],
            ["tools/hey-wrapper.sh", "Load test para demonstrar escalabilidade do HPA"],
            ["destroy-all.sh", "Destroi recursos para limpar ambiente"],
        ],
    )
    add_body(doc, "Sequência do run-all.sh (9 passos):", bold=True)
    for i, step in enumerate([
        "Verificar Helm instalado",
        "00-check-account.sh — atualizar LabRole ARN",
        "terraform init + apply — infraestrutura completa",
        "aws eks update-kubeconfig — configurar kubectl",
        "build-and-push.sh — imagens no ECR",
        "Metrics Server + Nginx Ingress (Helm)",
        "generate-secrets.sh — Secrets a partir do Terraform",
        "deploy-k8s.sh — manifestos Kubernetes",
        "run-migrations.sh + update-api-key.sh + teste de fluxo",
    ], 1):
        add_bullet(doc, f"Passo {i}: {step}")

    add_page_break(doc)

    # ── 6. Cluster K8s ──
    add_heading(doc, "6. Configuração do Cluster Kubernetes", level=1)
    add_body(doc,
        "Após o terraform apply, o cluster EKS togglemaster-cluster é configurado via kubectl. "
        "Componentes instalados pelos scripts (não pelo Terraform):"
    )
    add_heading(doc, "6.1 Metrics Server", level=2)
    add_body(doc,
        "Instalado via kubectl apply — necessário para o HPA monitorar utilização de CPU."
    )
    add_heading(doc, "6.2 Nginx Ingress Controller", level=2)
    add_body(doc,
        "Instalado via Helm (ingress-nginx). Com a LabRole nos nodes, o controller provisiona "
        "automaticamente um Load Balancer (ALB/NLB) na AWS. Rotas: /auth, /flags, /targeting, "
        "/evaluate, /analytics."
    )
    add_heading(doc, "6.3 Conexão ao cluster", level=2)
    add_body(doc, "aws eks update-kubeconfig --region us-east-1 --name togglemaster-cluster")

    # ── 7. Manifestos ──
    add_heading(doc, "7. Orquestração e Implantação (Manifestos)", level=1)
    add_body(doc, "Manifestos em infra/k8s/, aplicados por deploy-k8s.sh:")
    add_table(doc,
        ["Recurso K8s", "Conteúdo"],
        [
            ["namespace.yaml", "Namespace togglemaster"],
            ["secrets.yaml", "Gerado automaticamente por generate-secrets.sh (base64)"],
            ["configmap.yaml", "URLs internas, AWS_REGION, AWS_DYNAMODB_TABLE"],
            ["*-service.yaml (×5)", "Deployment + Service ClusterIP, requests/limits, liveness/readiness probes"],
            ["ingress.yaml", "Roteamento externo via Nginx com rewrite"],
            ["hpa.yaml", "HPA evaluation-service (2–8) e analytics-service (1–4), CPU 70%"],
        ],
    )
    add_body(doc, "Boas práticas aplicadas:", bold=True)
    add_bullet(doc, "Requests e Limits de CPU/memória em todos os Deployments;")
    add_bullet(doc, "LivenessProbe e ReadinessProbe;")
    add_bullet(doc, "Secrets em base64, gerados a partir dos outputs do Terraform;")
    add_bullet(doc, "Imagens do ECR com tag latest, buildadas para linux/amd64.")

    # ── 8. HPA ──
    add_heading(doc, "8. Escalabilidade (HPA)", level=1)
    add_table(doc,
        ["Serviço", "Min", "Max", "Métrica", "Target"],
        [
            ["evaluation-service", "2", "8", "CPU", "70%"],
            ["analytics-service", "1", "4", "CPU", "70%"],
        ],
    )
    add_body(doc,
        "Utilizamos HPA por CPU (Opção A — AWS Academy). O load test com hey "
        "(tools/hey-wrapper.sh, 150 conexões por 3 minutos) gera carga no /evaluate e "
        "demonstra o HPA escalando réplicas. Para analytics-service, mensagens enviadas "
        "manualmente à fila SQS aumentam CPU do worker e disparam scale-out."
    )
    add_body(doc,
        "Justificativa: KEDA (escalonamento por queueDepth) requer IRSA e novas IAM roles, "
        "indisponíveis no AWS Academy. HPA por CPU atende ao requisito mínimo da Fase 02."
    )

    add_page_break(doc)

    # ── 9. Data Stores ──
    add_heading(doc, "9. Arquitetura e Data Stores", level=1)
    add_heading(doc, "9.1 RDS (PostgreSQL) — dados transacionais", level=2)
    add_body(doc,
        "3 instâncias independentes (auth_db, flag_db, targeting_db). Dados estruturados de "
        "configuração com consistência ACID. Subnets privadas, acesso via Security Group dos nodes EKS."
    )
    add_heading(doc, "9.2 ElastiCache (Redis) — cache de alta performance", level=2)
    add_body(doc,
        "Cache efêmero (TTL 30s) no evaluation-service. Reduz latência do hot path evitando "
        "consultas síncronas a flag-service e targeting-service a cada requisição."
    )
    add_heading(doc, "9.3 DynamoDB — analytics de eventos", level=2)
    add_body(doc,
        "Tabela analytics-events (PK: event_id). analytics-service consome SQS e persiste "
        "user_id, flag_name, result e timestamp. PAY_PER_REQUEST adequado ao Academy."
    )
    add_heading(doc, "9.4 SQS — desacoplamento assíncrono", level=2)
    add_body(doc,
        "Fila togglemaster-events conecta evaluation-service (produtor) e analytics-service "
        "(consumidor), permitindo resposta imediata ao cliente sem aguardar persistência."
    )

    # ── 10. Segurança ──
    add_heading(doc, "10. Segurança", level=1)
    add_table(doc,
        ["Medida", "Implementação"],
        [
            ["Isolamento de rede", "RDS, Redis e nodes EKS em subnets privadas + NAT Gateway"],
            ["Security Groups", "RDS (5432) e Redis (6379) acessíveis apenas pelo SG dos nodes"],
            ["LabRole via Terraform", "Cluster e nodes usam LabRole existente (sem novas IAM roles)"],
            ["Autenticação API", "Chaves hasheadas SHA-256 no auth-service"],
            ["Secrets automatizados", "generate-secrets.sh lê terraform output, sem credenciais hardcoded no YAML"],
            ["IMDS hop limit", "Launch Template com hop_limit=2 para credenciais AWS nos pods"],
            ["ECR scan", "scan_on_push habilitado nos repositórios"],
        ],
    )

    # ── 11. Desafios ──
    add_heading(doc, "11. Desafios Enfrentados", level=1)
    desafios = [
        ("LabRole muda a cada sessão Academy",
         "O Account ID e ARN da LabRole mudam quando se abre nova sessão. Solução: script "
         "00-check-account.sh atualiza terraform.tfvars automaticamente antes do apply."),
        ("IMDS hop limit e credenciais AWS nos pods",
         "Pods não conseguiam acessar SQS/DynamoDB (NoCredentialProviders). Solução: Launch "
         "Template com http_put_response_hop_limit=2 no eks.tf."),
        ("Nodes em subnet privada sem internet",
         "Pull de imagens e APIs AWS falhavam. Solução: NAT Gateway + route table privada (nat.tf)."),
        ("DB Subnet Group",
         "AWS exige subnets em 2 AZs. Configurado private_a e private_b no rds.tf."),
        ("Geração automática de Secrets",
         "Endpoints RDS mudam a cada apply. Solução: generate-secrets.sh lê terraform output e "
         "regenera infra/k8s/secrets.yaml."),
        ("SERVICE_API_KEY dinâmica",
         "evaluation-service precisa de chave válida. Solução: update-api-key.sh cria key via "
         "auth-service e faz patch no Secret + rollout restart."),
        ("Build ARM vs AMD64",
         "Imagens buildadas em Mac M1 não rodavam nos nodes Linux. Solução: docker build "
         "--platform linux/amd64 no build-and-push.sh."),
        ("LocalStack timing local",
         "Container localstack-setup garante fila SQS pronta antes dos serviços iniciarem."),
    ]
    for titulo, desc in desafios:
        add_body(doc, titulo, bold=True)
        add_body(doc, desc)

    add_page_break(doc)

    # ── 12. Conclusão ──
    add_heading(doc, "12. Conclusão", level=1)
    add_body(doc,
        "A Fase 02 do ToggleMaster foi concluída com infraestrutura 100% automatizada via "
        "Terraform e pipeline de deploy via scripts Bash. Um único terraform apply provisiona "
        "VPC, NAT, EKS, node group, RDS, ElastiCache, SQS, DynamoDB e ECR."
    )
    add_body(doc, "Entregáveis atendidos:", bold=True)
    add_bullet(doc, "5 Dockerfiles multi-stage + docker-compose.yml funcional;")
    add_bullet(doc, "Infraestrutura completa via Terraform (incluindo EKS com LabRole);")
    add_bullet(doc, "Scripts de automação (run-all.sh) para deploy end-to-end;")
    add_bullet(doc, "Metrics Server + Nginx Ingress + manifestos K8s completos;")
    add_bullet(doc, "HPA para evaluation-service e analytics-service;")
    add_bullet(doc, "Demonstração de escalabilidade (hey + SQS) e dados no DynamoDB;")
    add_bullet(doc, "Vídeo de demonstração (link na seção Links do Projeto).")

    # ── 13. Referências ──
    add_heading(doc, "13. Referências", level=1)
    refs = [
        f"Repositório: {REPO_URL} (branch feature/melhorias-scripts)",
        "Enunciado Fase 02: POSTECH - Tech Challenge - Fase 2.pdf",
        "Terraform AWS Provider: https://registry.terraform.io/providers/hashicorp/aws",
        "Código IaC: pasta infra/ (vpc.tf, nat.tf, eks.tf, rds.tf, elasticache.tf, sqs.tf, dynamodb.tf, ecr.tf)",
        "Scripts de automação: run-all.sh, build-and-push.sh, generate-secrets.sh, deploy-k8s.sh",
        "Kubernetes Metrics Server: https://github.com/kubernetes-sigs/metrics-server",
        "Nginx Ingress: https://kubernetes.github.io/ingress-nginx/",
        f"Vídeo de demonstração: {VIDEO_URL}",
    ]
    for i, ref in enumerate(refs, 1):
        add_body(doc, f"{i}. {ref}")

    doc.save(OUTPUT)
    print(f"Relatorio gerado: {OUTPUT}")


if __name__ == "__main__":
    build_document()
