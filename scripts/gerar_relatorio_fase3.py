"""Gera o Relatório de Entrega da Fase 3 do ToggleMaster em .docx.

Uso: python scripts/gerar_relatorio_fase3.py
Requer: pip install python-docx

Campos a preencher antes de entregar: link do vídeo e print da estimativa de custos
(AWS Pricing Calculator) — marcados com [PREENCHER].
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                      "Relatorio_de_Entrega_Fase3.docx")

REPO_URL = "https://github.com/Maic2018/tech_challenge_3"
VIDEO_URL = "[PREENCHER: link do vídeo no YouTube, até 20 min]"

PARTICIPANTES = [
    ("Juliette Engel", "RM 373836"),
    ("Matheus Yuri Rodrigues da Silva", "RM 370672"),
    ("Michael Jonathan Venute Viana", "RM 371791"),
    ("Perla Coutinho Barbosa", "RM 371384"),
    ("Rodrigo de Albuquerque Andrade", "RM 373167"),
]


def spacing(paragraph, before=0, after=6, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    spacing(p, after=8)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    spacing(p, after=4)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return t


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── Capa ────────────────────────────────────────────────────────────────
    for text, size, bold in (("TECH CHALLENGE", 24, True), ("ToggleMaster", 20, True),
                             ("FASE 03 — IaC, DevSecOps e GitOps", 16, False),
                             ("DevOps e Arquitetura Cloud", 14, False)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
    doc.add_paragraph()
    body(doc, "Nome dos participantes – Grupo 45:", bold=True)
    for nome, rm in PARTICIPANTES:
        bullet(doc, f"{nome} – {rm}")
    body(doc, "Data: [PREENCHER]")
    doc.add_page_break()

    # ── Links ───────────────────────────────────────────────────────────────
    heading(doc, "Links do projeto")
    bullet(doc, f"Repositório GitHub (código, Terraform, workflows, GitOps): {REPO_URL}")
    bullet(doc, f"Documentação: {REPO_URL}/blob/main/README.md")
    bullet(doc, f"Vídeo de demonstração: {VIDEO_URL}")

    # ── Resumo ──────────────────────────────────────────────────────────────
    heading(doc, "Resumo")
    body(doc, (
        "Na Fase 3 a operação do ToggleMaster foi automatizada de ponta a ponta seguindo a regra "
        "\"se não está no código, não existe\". Toda a infraestrutura AWS (VPC, EKS com a LabRole, três "
        "RDS PostgreSQL, ElastiCache Redis, DynamoDB, SQS e cinco repositórios ECR) é criada por Terraform "
        "modularizado, com state remoto em S3 versionado e lock nativo (use_lockfile). Cada microsserviço "
        "tem um pipeline de CI no GitHub Actions com build e testes, lint, SCA e SAST com regra de bloqueio "
        "para vulnerabilidades críticas, build da imagem, scan de container com Trivy e push para o ECR com "
        "a tag do commit. O deploy é GitOps: o pipeline atualiza a tag no deployment.yaml da pasta gitops/ e o "
        "ArgoCD, instalado via Terraform (provider helm), sincroniza automaticamente os cinco microsserviços "
        "no cluster. Credenciais deixaram de existir em arquivos de texto: são geradas pelo Terraform, "
        "guardadas no Secrets Manager e entregues ao cluster como Secrets."
    ))

    # ── IaC ─────────────────────────────────────────────────────────────────
    heading(doc, "1. Infraestrutura como Código (Terraform)")
    table(doc, ["Módulo", "Recursos"], [
        ("modules/network", "VPC 10.0.0.0/16, 2 subnets públicas + 2 privadas, IGW, NAT Gateway, route tables"),
        ("modules/eks", "Cluster EKS 1.33 com LabRole (data source), Managed Node Group t3.medium (1/2/4), "
                        "launch template com IMDS hop limit 2, security group dos nodes"),
        ("modules/rds", "3 instâncias PostgreSQL 15 (auth_db, flag_db, targeting_db), criptografadas, privadas"),
        ("modules/elasticache", "Cluster Redis 7 (cache.t3.micro)"),
        ("modules/dynamodb", "Tabela ToggleMasterAnalytics (PAY_PER_REQUEST)"),
        ("modules/sqs", "Fila togglemaster-events + dead-letter queue"),
        ("modules/ecr", "5 repositórios com scan on push e lifecycle policy"),
        ("infra/secrets.tf", "random_password do RDS + segredo no AWS Secrets Manager"),
        ("infra/platform", "ArgoCD, ingress-nginx, metrics-server (Helm), namespace e Secrets do Kubernetes"),
    ])
    body(doc, (
        "Estado remoto: backend S3 com encrypt e use_lockfile, bucket criado por scripts/bootstrap-backend.sh "
        "(versionamento, SSE-S3 e bloqueio de acesso público). O nome do bucket inclui o Account ID, que muda "
        "entre sessões do AWS Academy, e é informado no init pelo script scripts/tf-init.sh."
    ))

    # ── CI ──────────────────────────────────────────────────────────────────
    heading(doc, "2. Pipeline de CI e DevSecOps (GitHub Actions)")
    table(doc, ["Job", "Ferramentas", "Regra de bloqueio"], [
        ("1. Build & Unit Test", "go build / go test -race; pip install / compileall / pytest", "falha de build ou teste"),
        ("2. Linter", "golangci-lint; flake8", "qualquer achado"),
        ("3. Security Scan (SCA)", "Trivy fs (go.mod/go.sum, requirements.txt) + detecção de segredos", "CRITICAL"),
        ("3. Security Scan (SAST)", "gosec; bandit", "HIGH (gosec) / MEDIUM+ (bandit)"),
        ("4. Docker Build & Push", "docker build, Trivy image, login ECR, push v1.0.0-<sha>", "CRITICAL na imagem"),
        ("5. GitOps", "scripts/update-image-tag.sh + commit [skip ci]", "—"),
    ])
    body(doc, (
        "Um workflow reutilizável (_service-ci.yml) concentra a lógica; cinco workflows finos "
        "(auth-service.yml, flag-service.yml, targeting-service.yml, evaluation-service.yml, "
        "analytics-service.yml) disparam em Pull Request e push na main quando a pasta do serviço muda."
    ))

    # ── GitOps ──────────────────────────────────────────────────────────────
    heading(doc, "3. Entrega Contínua e GitOps (ArgoCD)")
    bullet(doc, "Repositório GitOps: pasta gitops/ do monorepo, apenas manifestos (Deployment, Service, HPA, ConfigMap, Ingress).")
    bullet(doc, "ArgoCD instalado por Terraform (helm_release) com UI exposta por Load Balancer.")
    bullet(doc, "ApplicationSet com generator git directories gera uma Application por microsserviço; "
                "uma Application adicional cuida do ConfigMap e do Ingress.")
    bullet(doc, "Sync automático com prune e selfHeal; reconciliação a cada 60 segundos.")
    bullet(doc, "Ao final do CI, o job 5 altera a linha image: do deployment.yaml e faz o commit; o ArgoCD detecta e sincroniza.")

    # ── Desafios ────────────────────────────────────────────────────────────
    heading(doc, "4. Desafios encontrados e decisões tomadas")
    for txt in (
        "AWS Academy não permite criar roles: a LabRole é lida via data source e o pipeline autentica no ECR "
        "com access keys da sessão do laboratório publicadas como Secrets do GitHub (renovadas por script), "
        "já que não é possível criar uma role OIDC para o GitHub Actions.",
        "Account ID muda entre sessões: nomes de bucket e registry são derivados em tempo de execução; o CI "
        "grava no GitOps a imagem completa do registry em que fez login, então os manifestos se corrigem sozinhos.",
        "Segredos em texto plano (Fase 2): substituídos por random_password no Terraform, Secrets Manager e "
        "Secrets do Kubernetes criados por Terraform. O state e o tfvars saíram do repositório e o histórico "
        "deve ser limpo com git filter-repo; a senha antiga foi descartada com a infra.",
        "Dependências vulneráveis: as versões antigas (pgx 5.5.0, Flask 2.2, Python 3.9, Go 1.21, Alpine 3.19) "
        "acusavam CRITICAL no Trivy. Foram atualizadas e as imagens migraram para distroless (Go) e "
        "python:3.12-slim não-root.",
        "Cinco pipelines commitando no mesmo repositório: um concurrency group serializa o job de GitOps e "
        "um loop de retry com rebase evita conflitos; [skip ci] evita loops de execução.",
        "Chicken-and-egg da chave interna do evaluation-service: a chave é gerada pelo Terraform e o hash "
        "SHA-256 é inserido no banco pelo Job de migrations, sem edição manual de Secret.",
        "Ordem de destruição: o destroy da plataforma remove os Load Balancers criados pelo Helm antes do "
        "destroy da VPC, evitando o erro DependencyViolation.",
    ):
        bullet(doc, txt)

    # ── Custos ──────────────────────────────────────────────────────────────
    heading(doc, "5. Estimativa de custos AWS")
    body(doc, "[PREENCHER: inserir aqui o print da estimativa do AWS Pricing Calculator]", bold=True)
    body(doc, "Referência mensal aproximada (us-east-1, uso contínuo), para conferir com a calculadora:")
    table(doc, ["Recurso", "Dimensionamento", "Estimativa/mês (USD)"], [
        ("EKS control plane", "1 cluster", "73"),
        ("EC2 nodes", "2× t3.medium on-demand", "60"),
        ("RDS PostgreSQL", "3× db.t3.micro, 20 GB gp2 cada", "45"),
        ("ElastiCache Redis", "1× cache.t3.micro", "12"),
        ("NAT Gateway", "1 + tráfego", "35"),
        ("Load Balancers", "2× Classic (ingress-nginx, ArgoCD)", "36"),
        ("DynamoDB / SQS / ECR / S3 / Secrets Manager", "on-demand, baixo volume", "5"),
        ("Total aproximado", "", "≈ 265"),
    ])
    body(doc, "No AWS Academy o ambiente é criado apenas durante as sessões e destruído com destroy-all.sh, "
              "então o custo real fica bem abaixo do valor mensal.")

    # ── Conclusão ───────────────────────────────────────────────────────────
    heading(doc, "6. Conclusão")
    body(doc, (
        "A Fase 3 entrega infraestrutura imutável e reproduzível (um único terraform apply por root module), "
        "pipelines que bloqueiam vulnerabilidades críticas antes da imagem chegar ao registry e deploy "
        "declarativo por GitOps, em que o cluster converge para o que está versionado. Recriar o ambiente de "
        "homologação passou de dias para cerca de 30 minutos com o run-all.sh."
    ))

    doc.save(OUTPUT)
    print(f"Relatório gerado em: {OUTPUT}")


if __name__ == "__main__":
    build()
